"""Renderer produces a valid .docx with the template's sections + placeholders."""

from __future__ import annotations

from docx import Document

from asm_review.render.docx_renderer import save_review_note


def _all_text(path) -> str:
    doc = Document(str(path))
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


def test_renders_sections_values_and_placeholders(sample_note, tmp_path) -> None:
    out = tmp_path / "note.docx"
    save_review_note(sample_note, out)
    assert out.exists()

    text = _all_text(out)
    # Title + a few section headings
    assert "ASM REVIEW NOTE" in text
    assert "AUDIT DETAILS" in text
    assert "DETAILS OF BANKING EXPOSURE OF THE CUSTOMER" in text
    assert "REMARKS ON ASM REPORT" in text
    # An extracted value
    assert "Acme Industries Pvt Ltd" in text
    # A missing field -> placeholder (external_rating was not found)
    assert "[To be entered by L1]" in text
    # Table-heavy doc
    assert len(Document(str(out)).tables) >= 10


def test_custom_placeholder(sample_note, tmp_path) -> None:
    out = tmp_path / "note.docx"
    save_review_note(sample_note, out, placeholder="<<FILL>>")
    text = _all_text(out)
    assert "<<FILL>>" in text
    assert "[To be entered by L1]" not in text
