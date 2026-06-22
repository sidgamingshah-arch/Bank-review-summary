"""ReviewNote -> .docx, matching the pasted ASM review-note template.

Layout/labels are modelled on the pasted format and are the part most likely to
need a refinement pass once a real .docx template is available (exact fonts,
column widths, numbering). Any field that was not found renders as the configured
placeholder (default ``[To be entered by L1]``).
"""

from __future__ import annotations

from pathlib import Path
from typing import Optional, Sequence

from docx import Document
from docx.document import Document as DocxDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from asm_review.schema.fields import Field
from asm_review.schema.models import ReviewNote

DEFAULT_PLACEHOLDER = "[To be entered by L1]"


class DocxRenderer:
    def __init__(self, placeholder: str = DEFAULT_PLACEHOLDER) -> None:
        self.placeholder = placeholder

    # --- small helpers ----------------------------------------------------
    def _v(self, field: Field) -> str:
        disp = field.display()
        return disp if disp is not None else self.placeholder

    @staticmethod
    def _set_cell(cell, text: Optional[str], *, bold: bool = False) -> None:
        cell.text = text or ""
        para = cell.paragraphs[0]
        if not para.runs:
            para.add_run("")
        para.runs[0].bold = bold

    @staticmethod
    def _heading(doc: DocxDocument, text: str, *, size: int = 12) -> None:
        doc.add_paragraph()
        para = doc.add_paragraph()
        run = para.add_run(text)
        run.bold = True
        run.font.size = Pt(size)

    @staticmethod
    def _grid(doc: DocxDocument, rows: int, cols: int):
        table = doc.add_table(rows=rows, cols=cols)
        table.style = "Table Grid"
        return table

    def _kv(self, doc: DocxDocument, pairs: Sequence[tuple[str, Field]]) -> None:
        table = self._grid(doc, len(pairs), 2)
        table.columns[0].width = Pt(200)
        for row, (label, field) in zip(table.rows, pairs):
            self._set_cell(row.cells[0], label, bold=True)
            self._set_cell(row.cells[1], self._v(field))

    def _numbered(self, doc: DocxDocument, items: Sequence[tuple[str, str, Field]]) -> None:
        table = self._grid(doc, len(items), 3)
        for row, (no, label, field) in zip(table.rows, items):
            self._set_cell(row.cells[0], no, bold=True)
            self._set_cell(row.cells[1], label)
            self._set_cell(row.cells[2], self._v(field))

    def _data_table(
        self, doc: DocxDocument, headers: Sequence[str], rows: Sequence[Sequence[str]]
    ) -> None:
        body = rows or [[self.placeholder] * len(headers)]
        table = self._grid(doc, len(body) + 1, len(headers))
        for cell, head in zip(table.rows[0].cells, headers):
            self._set_cell(cell, head, bold=True)
        for trow, cells in zip(table.rows[1:], body):
            for cell, text in zip(trow.cells, cells):
                self._set_cell(cell, text)

    # --- top-level render -------------------------------------------------
    def render(self, note: ReviewNote) -> DocxDocument:
        doc = Document()
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        trun = title.add_run("ASM REVIEW NOTE")
        trun.bold = True
        trun.font.size = Pt(14)
        doc.add_paragraph(f"Date : {self._v(note.header.date_of_review)}")

        self._identifiers(doc, note)
        self._audit_details(doc, note)
        self._account_details(doc, note)
        self._banking_exposure(doc, note)
        self._bank_exposure(doc, note)
        self._primary_security(doc, note)
        self._main_observations(doc, note)
        self._insurance(doc, note)
        self._project(doc, note)
        self._drawing_power(doc, note)
        self._account_position(doc, note)
        self._crilc_mca(doc, note)
        self._remarks(doc, note)
        return doc

    # --- sections ---------------------------------------------------------
    def _identifiers(self, doc: DocxDocument, note: ReviewNote) -> None:
        h = note.header
        table = self._grid(doc, 4, 4)
        self._set_cell(table.rows[0].cells[0], "Br Code", bold=True)
        self._set_cell(table.rows[0].cells[1], self._v(h.br_code))
        self._set_cell(table.rows[0].cells[2], "Br Name", bold=True)
        self._set_cell(table.rows[0].cells[3], self._v(h.br_name))
        self._set_cell(table.rows[1].cells[0], "RO", bold=True)
        self._set_cell(table.rows[1].cells[1], self._v(h.ro))
        self._set_cell(table.rows[2].cells[0], "CIF ID", bold=True)
        self._set_cell(table.rows[2].cells[1], self._v(h.cif_id))
        self._set_cell(table.rows[3].cells[0], "A/c Name", bold=True)
        self._set_cell(table.rows[3].cells[1], self._v(h.ac_name))

    def _audit_details(self, doc: DocxDocument, note: ReviewNote) -> None:
        self._heading(doc, "AUDIT DETAILS")
        a = note.audit_details
        self._numbered(
            doc,
            [
                ("1", "Name of ASM Auditor", a.name_of_asm_auditor),
                ("2", "Appointed by", a.appointed_by),
                ("3", "Audit as on quarter", a.audit_as_on_quarter),
                ("4", "Report Date", a.report_date),
            ],
        )

    def _account_details(self, doc: DocxDocument, note: ReviewNote) -> None:
        self._heading(doc, "ACCOUNT DETAILS")
        a = note.account_details
        self._numbered(
            doc,
            [
                ("1", "Constitution", a.constitution),
                ("2", "Nature of Activity", a.nature_of_activity),
                ("3", "Banking arrangement", a.banking_arrangement),
                ("4", "BANK Sanction Ref", a.bank_sanction_ref),
                ("5", "Business Vertical", a.business_vertical),
                ("6", "BANK Rating", a.bank_rating),
                ("7", "External Rating", a.external_rating),
            ],
        )

    def _banking_exposure(self, doc: DocxDocument, note: ReviewNote) -> None:
        self._heading(doc, "I.  DETAILS OF BANKING EXPOSURE OF THE CUSTOMER")
        be = note.banking_exposure
        rows: list[list[str]] = []
        for i, r in enumerate(be.rows, start=1):
            rows.append(
                [
                    str(i),
                    self._v(r.lender),
                    self._v(r.fbwc_os_balance),
                    self._v(r.nfbwc_os_balance),
                    self._v(r.term_loan),
                ]
            )
        rows.append(
            [
                "",
                "TOTAL",
                self._v(be.total_fbwc),
                self._v(be.total_nfbwc),
                self._v(be.total_term_loan),
            ]
        )
        self._data_table(
            doc,
            ["Sl No", "Lenders", "FBWC o/s balance", "NFBWC o/s b/s", "TL"],
            rows,
        )

    def _bank_exposure(self, doc: DocxDocument, note: ReviewNote) -> None:
        self._heading(doc, "II.  DETAILS OF EXPOSURE WITH BANK")
        rows = [
            [self._v(r.facility), self._v(r.limit_crs), self._v(r.dp_crs), self._v(r.balance_crs)]
            for r in note.bank_exposure.rows
        ]
        self._data_table(
            doc, ["Facility", "Limit (crs)", "DP as on (Crs)", "Bal as on (Crs)"], rows
        )
        if not note.bank_exposure.footnote.needs_manual_entry:
            doc.add_paragraph(f"* {self._v(note.bank_exposure.footnote)}")

    def _primary_security(self, doc: DocxDocument, note: ReviewNote) -> None:
        self._heading(
            doc, "III.  DETAILS OF PRIMARY SECURITY WITH BANK AS PER SANCTION TERMS"
        )
        ps = note.primary_security
        doc.add_paragraph(f"Primary Security : {self._v(ps.primary_security)}")
        doc.add_paragraph(f"Margin : {self._v(ps.margin)}")

    def _main_observations(self, doc: DocxDocument, note: ReviewNote) -> None:
        self._heading(doc, "IV.  MAIN OBSERVATIONS IN ASM REPORT")
        m = note.main_observations
        self._kv(
            doc,
            [
                ("Business Positions", m.business_positions),
                ("Networth level", m.networth_level),
                ("Financial Ratios", m.financial_ratios),
                ("Sales and Purchase and Profit levels", m.sales_purchase_profit_levels),
                ("Asset and liabilities", m.asset_and_liabilities),
                ("Contingent & statutory liabilities", m.contingent_statutory_liabilities),
                ("Sundry Creditors", m.sundry_creditors),
                ("Sundry Debtors / book debts", m.sundry_debtors_book_debts),
                ("Stock / Raw materials / WIP", m.stock_raw_materials_wip),
                ("Fixed assets", m.fixed_assets),
                (
                    "Bank Borrowings / limit utilisations / overdues etc",
                    m.bank_borrowings_limit_utilisations_overdues,
                ),
                ("Bank account operations", m.bank_account_operations),
                ("Cash flow and ALM", m.cash_flow_and_alm),
                (
                    "High value transactions with sister/associate concern & related parties",
                    m.high_value_related_party_transactions,
                ),
                ("Any other observations by auditor", m.any_other_observations),
            ],
        )

    def _insurance(self, doc: DocxDocument, note: ReviewNote) -> None:
        self._heading(doc, "V.  Insurance")
        ins = note.insurance
        self._kv(
            doc,
            [
                ("Whether adequate insurance present", ins.whether_adequate_insurance_present),
                ("Validity of insurance", ins.validity_of_insurance),
                ("Whether Bank's Lien noted", ins.whether_banks_lien_noted),
                ("Whether all locations covered", ins.whether_all_locations_covered),
                ("Whether all risks are covered", ins.whether_all_risks_covered),
                ("Other Remarks", ins.other_remarks),
            ],
        )

    def _project(self, doc: DocxDocument, note: ReviewNote) -> None:
        self._heading(doc, "VI.  OBSERVATIONS ON PROJECT UNDER IMPLEMENTATION")
        p = note.project_observations
        self._kv(
            doc,
            [
                (
                    "Deviations in project progress vis-a-vis timelines and amount disbursed",
                    p.deviations_in_project_progress,
                ),
                (
                    "Scheduled DCCO and number of months remaining to DCCO",
                    p.scheduled_dcco_and_months_remaining,
                ),
                ("Any other observations by the auditor", p.any_other_observations),
            ],
        )

    def _drawing_power(self, doc: DocxDocument, note: ReviewNote) -> None:
        self._heading(doc, "VII.  DRAWING POWER COMPUTATION")
        rows = [[self._v(r.particulars), self._v(r.value)] for r in note.drawing_power.rows]
        self._data_table(doc, ["Particulars", "Value"], rows)

    def _account_position(self, doc: DocxDocument, note: ReviewNote) -> None:
        self._heading(doc, "VIII.  POSITION OF ACCOUNT AS ON DATE OF AUDIT")
        rows = [
            [
                self._v(r.nature),
                self._v(r.total_limit),
                self._v(r.total_dp),
                self._v(r.total_balance_os),
                self._v(r.bank_limit),
                self._v(r.bank_dp),
                self._v(r.bank_balance_os),
                self._v(r.remarks),
                self._v(r.deviation_pct),
            ]
            for r in note.account_position.rows
        ]
        self._data_table(
            doc,
            [
                "Nature",
                "Total Limit (Mul/Con)",
                "Total DP (Mul/Con)",
                "Total bal o/s as on audit date",
                "BANK limit",
                "BANK DP",
                "Bal O/s",
                "Remarks",
                "Dev %",
            ],
            rows,
        )
        doc.add_paragraph(
            "* DP details will be updated with stock audit team under LAM cell."
        )

    def _crilc_mca(self, doc: DocxDocument, note: ReviewNote) -> None:
        self._heading(doc, "IX.  DETAILS OF VERIFICATION DONE FROM CRILC AND MCA DATA")
        para = doc.add_paragraph()
        para.add_run("CRILC DATA").bold = True
        c = note.crilc_mca
        self._numbered(
            doc,
            [
                (
                    "1",
                    "Whether the party is maintaining current accounts with other bank/s",
                    c.maintaining_current_accounts_other_banks,
                ),
                ("1(a)", "If yes, name of bank/s", c.names_of_banks),
                ("1(b)", "Whether permission obtained", c.whether_permission_obtained),
                ("2", "Whether ROC charges are filed properly", c.whether_roc_charges_filed_properly),
            ],
        )

    def _remarks(self, doc: DocxDocument, note: ReviewNote) -> None:
        self._heading(doc, "X.  REMARKS ON ASM REPORT")
        doc.add_paragraph(
            "Comparison with previous quarter ASM report critical observations, if any"
        )
        rows = [
            [str(i), self._v(r.previous_quarter_comment), self._v(r.current_quarter_status)]
            for i, r in enumerate(note.remarks.prev_quarter_comparison, start=1)
        ]
        self._data_table(
            doc,
            [
                "Sl No",
                "Comment in previous quarter ASM report",
                "Status as per current quarter of audit",
            ],
            rows,
        )

        def labelled(label: str, field: Field) -> None:
            p = doc.add_paragraph()
            p.add_run(label).bold = True
            doc.add_paragraph(self._v(field))

        labelled("Critical Observations", note.remarks.critical_observations)
        labelled("Other Observations", note.remarks.other_observations)
        labelled(
            "For information of sanctioning authority / business vertical",
            note.remarks.for_sanctioning_authority,
        )
        labelled("For information of reviewing authority", note.remarks.for_reviewing_authority)


def render_review_note(note: ReviewNote, placeholder: str = DEFAULT_PLACEHOLDER) -> DocxDocument:
    return DocxRenderer(placeholder).render(note)


def save_review_note(
    note: ReviewNote, path: str | Path, placeholder: str = DEFAULT_PLACEHOLDER
) -> Path:
    doc = render_review_note(note, placeholder)
    path = Path(path)
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(path))
    return path
