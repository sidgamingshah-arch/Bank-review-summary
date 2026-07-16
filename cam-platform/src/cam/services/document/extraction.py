"""Per-format text extraction (FR-C05).

PDF extraction reads the text layer only — a scanned/image-only PDF yields no
text and the document is flagged ``no_text`` (OCR is a documented v1 gap).
Extract size is capped so a rogue file cannot bloat the extract store.
"""
from __future__ import annotations

import io

import docx
import openpyxl
from pypdf import PdfReader

MAX_EXTRACT_CHARS = 200_000
MAX_XLSX_ROWS_PER_SHEET = 200


def _pdf(content: bytes) -> str:
    reader = PdfReader(io.BytesIO(content))
    return "\n".join((page.extract_text() or "") for page in reader.pages)


def _docx(content: bytes) -> str:
    document = docx.Document(io.BytesIO(content))
    parts = [p.text for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text for cell in row.cells))
    return "\n".join(parts)


def _xlsx(content: bytes) -> str:
    workbook = openpyxl.load_workbook(io.BytesIO(content), read_only=True, data_only=True)
    parts: list[str] = []
    try:
        for sheet in workbook.worksheets:
            parts.append(f"[sheet: {sheet.title}]")
            for i, row in enumerate(sheet.iter_rows(values_only=True)):
                if i >= MAX_XLSX_ROWS_PER_SHEET:
                    parts.append("... (rows truncated)")
                    break
                parts.append(" | ".join("" if v is None else str(v) for v in row))
    finally:
        workbook.close()
    return "\n".join(parts)


def _plain(content: bytes) -> str:
    return content.decode("utf-8", errors="replace")


_EXTRACTORS = {".pdf": _pdf, ".docx": _docx, ".xlsx": _xlsx, ".csv": _plain, ".txt": _plain}


def extract_text(content: bytes, ext: str) -> str | None:
    """Extracted text capped at ``MAX_EXTRACT_CHARS``; ``None`` when the format
    is unsupported or the file cannot be parsed."""
    extractor = _EXTRACTORS.get((ext or "").lower())
    if extractor is None:
        return None
    try:
        return extractor(content)[:MAX_EXTRACT_CHARS]
    except Exception:
        return None
