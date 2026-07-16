"""DOCX / PDF exporters for CAM documents (FR-E07/E08).

Both renderers consume the full Cam JSON (as produced by the service, sections
in ``order`` with current head ``content``) and share ``cam.common.markdownish``
for the markdown-subset parsing, so the two formats stay in lockstep.

Watermark rule (FR-E08): while ``status == "draft"`` every export carries
"AI-ASSISTED DRAFT — not for credit decision" (page header + top banner);
finalisation drops it. The ``_gaps`` trailer is always rendered.
"""
from __future__ import annotations

import io

from docx import Document
from docx.shared import Pt, RGBColor
from fpdf import FPDF
from fpdf.enums import XPos, YPos

from cam.common.db import iso, utcnow
from cam.common.markdownish import Block, bold_spans, parse_blocks, strip_bold

WATERMARK = "AI-ASSISTED DRAFT — not for credit decision"
GAPS_INTRO = "Disclosed data gaps for this generation:"
AMBER = (0xB8, 0x86, 0x0B)  # amber-ish / dark goldenrod


# ---------------------------------------------------------------------------
# DOCX
# ---------------------------------------------------------------------------

def _docx_runs(paragraph, text: str, *, force_bold: bool = False) -> None:
    for span, bold in bold_spans(text):
        run = paragraph.add_run(span)
        run.bold = True if force_bold else bold


def _docx_table(doc, rows: list[list[str]]) -> None:
    cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=cols)
    table.style = "Table Grid"
    for r_idx, row in enumerate(rows):
        for c_idx in range(cols):
            text = strip_bold(row[c_idx]) if c_idx < len(row) else ""
            cell = table.cell(r_idx, c_idx)
            cell.text = ""
            run = cell.paragraphs[0].add_run(text)
            if r_idx == 0:
                run.bold = True  # header row


def _docx_block(doc, block: Block) -> None:
    if block.kind == "heading":
        doc.add_heading(strip_bold(block.text), level=min(block.level + 1, 4))
    elif block.kind == "bullets":
        for item in block.items:
            p = doc.add_paragraph(style="List Bullet")
            _docx_runs(p, item)
    elif block.kind == "table":
        _docx_table(doc, block.rows)
    else:  # paragraph
        p = doc.add_paragraph()
        _docx_runs(p, block.text)


def render_docx(cam: dict) -> bytes:
    doc = Document()
    draft = cam.get("status") == "draft"

    if draft:
        # FR-E08: watermark in the page header of every docx section ...
        for docx_section in doc.sections:
            header_p = docx_section.header.paragraphs[0]
            run = header_p.add_run(WATERMARK)
            run.bold = True
            run.font.color.rgb = RGBColor(*AMBER)
        # ... and a bold amber banner at the top of the body.
        banner = doc.add_paragraph()
        run = banner.add_run(WATERMARK)
        run.bold = True
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(*AMBER)

    doc.add_heading(cam.get("title", "Credit Assessment Memo"), level=0)
    meta = doc.add_paragraph()
    meta.add_run(
        f"Template: {cam.get('template_key', '')} · Status: {cam.get('status', '')} · "
        f"Generated: {iso(utcnow())} · Run: {cam.get('run_id', '')}"
    ).italic = True

    for section in cam.get("sections", []):
        doc.add_heading(section.get("name", section.get("section_code", "")), level=1)
        if section.get("section_code") == "_gaps":
            intro = doc.add_paragraph()
            intro.add_run(GAPS_INTRO).italic = True
        for block in parse_blocks(section.get("content") or ""):
            _docx_block(doc, block)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# PDF (fpdf2, core Helvetica fonts -> latin-1 text only)
# ---------------------------------------------------------------------------

_PDF_CHAR_MAP = {
    "₹": "Rs.", "€": "EUR",
    "–": "-", "—": "-", "−": "-",
    "‘": "'", "’": "'", "“": '"', "”": '"',
    "…": "...", " ": " ",
    "•": "\x95",  # bullet -> WinAnsi bullet glyph (renders as • in core fonts)
}


def pdf_text(text: str) -> str:
    """Sanitise arbitrary markdown text for fpdf's latin-1 core fonts."""
    for src, dst in _PDF_CHAR_MAP.items():
        text = text.replace(src, dst)
    return text.encode("latin-1", errors="replace").decode("latin-1")


class CamPdf(FPDF):
    def __init__(self, watermark: bool):
        super().__init__(format="A4")
        self.watermark = watermark
        self.set_auto_page_break(auto=True, margin=18)

    def header(self):  # every page (FR-E08)
        if not self.watermark:
            return
        self.set_font("Helvetica", "B", 9)
        self.set_text_color(*AMBER)
        self.cell(0, 6, pdf_text(WATERMARK), align="C")
        self.ln(9)
        self.set_text_color(0, 0, 0)

    def footer(self):
        self.set_y(-14)
        self.set_font("Helvetica", "I", 8)
        self.set_text_color(120, 120, 120)
        self.cell(0, 8, f"Page {self.page_no()}", align="C")
        self.set_text_color(0, 0, 0)


_HEADING_SIZES = {1: 16, 2: 13}  # level >= 3 -> 11


def _pdf_heading(pdf: FPDF, text: str, level: int) -> None:
    pdf.set_font("Helvetica", "B", _HEADING_SIZES.get(level, 11))
    pdf.ln(3)
    pdf.multi_cell(0, 7, pdf_text(strip_bold(text)), new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    pdf.ln(1)
    pdf.set_font("Helvetica", "", 10)


def _pdf_table(pdf: FPDF, rows: list[list[str]]) -> None:
    cols = max(len(r) for r in rows)
    padded = [[strip_bold(r[i]) if i < len(r) else "" for i in range(cols)] for r in rows]
    pdf.set_font("Helvetica", "", 9)
    if hasattr(pdf, "table"):  # fpdf2 >= 2.7 table() context manager
        with pdf.table() as table:
            for row in padded:
                table_row = table.row()
                for cell in row:
                    table_row.cell(pdf_text(cell))
    else:  # pragma: no cover - fallback for very old fpdf2
        width = pdf.epw / cols
        for r_idx, row in enumerate(padded):
            pdf.set_font("Helvetica", "B" if r_idx == 0 else "", 9)
            for cell in row:
                pdf.cell(width, 6, pdf_text(cell), border=1)
            pdf.ln(6)
    pdf.ln(2)
    pdf.set_font("Helvetica", "", 10)


def _pdf_block(pdf: FPDF, block: Block) -> None:
    if block.kind == "heading":
        _pdf_heading(pdf, block.text, min(block.level + 1, 4))
    elif block.kind == "bullets":
        pdf.set_font("Helvetica", "", 10)
        for item in block.items:
            pdf.multi_cell(0, 5.5, pdf_text("• " + strip_bold(item)),
                           new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        pdf.ln(1)
    elif block.kind == "table":
        _pdf_table(pdf, block.rows)
    else:
        pdf.set_font("Helvetica", "", 10)
        pdf.multi_cell(0, 5.5, pdf_text(strip_bold(block.text)), new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        pdf.ln(1)


def render_pdf(cam: dict) -> bytes:
    pdf = CamPdf(watermark=cam.get("status") == "draft")
    pdf.add_page()

    # Title block
    pdf.set_font("Helvetica", "B", 20)
    pdf.multi_cell(0, 10, pdf_text(cam.get("title", "Credit Assessment Memo")),
                   new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    pdf.set_font("Helvetica", "I", 9)
    pdf.set_text_color(90, 90, 90)
    pdf.multi_cell(0, 5, pdf_text(
        f"Template: {cam.get('template_key', '')} · Status: {cam.get('status', '')} · "
        f"Generated: {iso(utcnow())} · Run: {cam.get('run_id', '')}"
    ), new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    pdf.set_text_color(0, 0, 0)
    pdf.ln(4)

    for section in cam.get("sections", []):
        _pdf_heading(pdf, section.get("name", section.get("section_code", "")), 1)
        if section.get("section_code") == "_gaps":
            pdf.set_font("Helvetica", "I", 10)
            pdf.multi_cell(0, 5.5, pdf_text(GAPS_INTRO), new_x=XPos.LMARGIN, new_y=YPos.NEXT)
            pdf.ln(1)
            pdf.set_font("Helvetica", "", 10)
        for block in parse_blocks(section.get("content") or ""):
            _pdf_block(pdf, block)

    return bytes(pdf.output())
